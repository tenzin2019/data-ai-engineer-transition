"""
Document processing engine for extracting text from various document formats.
"""

import io
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import PyPDF2
import pdfplumber
import docx
import openpyxl
from PIL import Image
import pytesseract

# Try to import magic, with fallback
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    magic = None

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))

from models.document import DocumentType
from utils.file_utils import get_file_extension, validate_file_type

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing and text extraction from various formats."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_types = {
            'application/pdf': DocumentType.PDF,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentType.XLSX,
            'text/plain': DocumentType.TXT,
        }
    
    def process_document(self, file_path: Union[str, Path], mime_type: Optional[str] = None) -> Dict:
        """
        Process a document and extract text content.
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document (optional)
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect MIME type if not provided
        if not mime_type:
            if MAGIC_AVAILABLE:
                mime_type = magic.from_file(str(file_path), mime=True)
            else:
                # Fallback to file extension-based detection
                mime_type = self._get_mime_type_from_extension(file_path)
        
        # Validate file type
        if mime_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        document_type = self.supported_types[mime_type]
        
        try:
            # Extract text based on document type
            if document_type == DocumentType.PDF:
                return self._process_pdf(file_path)
            elif document_type == DocumentType.DOCX:
                return self._process_docx(file_path)
            elif document_type == DocumentType.XLSX:
                return self._process_xlsx(file_path)
            elif document_type == DocumentType.TXT:
                return self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported document type: {document_type}")
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """Process PDF documents and extract text."""
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                text_content = []
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
                
                extracted_text = "\n\n".join(text_content)
                
                # Extract metadata
                document_metadata = {
                    'page_count': page_count,
                    'creator': pdf.metadata.get('Creator', '') if pdf.metadata else '',
                    'producer': pdf.metadata.get('Producer', '') if pdf.metadata else '',
                    'creation_date': pdf.metadata.get('CreationDate', '') if pdf.metadata else '',
                    'modification_date': pdf.metadata.get('ModDate', '') if pdf.metadata else '',
                }
                
                return {
                    'text': extracted_text,
                    'page_count': page_count,
                    'metadata': document_metadata,
                    'document_type': DocumentType.PDF
                }
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
                
                extracted_text = "\n\n".join(text_content)
                
                # Extract metadata
                document_metadata = {}
                if pdf_reader.metadata:
                    document_metadata = {
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', '')),
                    }
                
                return {
                    'text': extracted_text,
                    'page_count': page_count,
                    'metadata': document_metadata,
                    'document_type': DocumentType.PDF
                }
    
    def _process_docx(self, file_path: Path) -> Dict:
        """Process DOCX documents and extract text."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = paragraphs + tables_text
            extracted_text = "\n\n".join(all_text)
            
            # Extract metadata
            core_props = doc.core_properties
            document_metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }
            
            return {
                'text': extracted_text,
                'page_count': len(paragraphs),  # Approximate page count
                'metadata': document_metadata,
                'document_type': DocumentType.DOCX
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise
    
    def _process_xlsx(self, file_path: Path) -> Dict:
        """Process XLSX documents and extract text."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            all_text = []
            sheet_count = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_count += 1
                
                # Add sheet header
                all_text.append(f"=== Sheet: {sheet_name} ===")
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    row_data = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                    
                    if row_data:
                        all_text.append(" | ".join(row_data))
                
                all_text.append("")  # Add spacing between sheets
            
            extracted_text = "\n".join(all_text)
            
            # Extract metadata
            document_metadata = {
                'sheet_count': sheet_count,
                'sheet_names': workbook.sheetnames,
                'creator': workbook.properties.creator or '',
                'created': str(workbook.properties.created) if workbook.properties.created else '',
                'modified': str(workbook.properties.modified) if workbook.properties.modified else '',
            }
            
            return {
                'text': extracted_text,
                'page_count': sheet_count,
                'metadata': document_metadata,
                'document_type': DocumentType.XLSX
            }
            
        except Exception as e:
            logger.error(f"Error processing XLSX file {file_path}: {str(e)}")
            raise
    
    def _process_txt(self, file_path: Path) -> Dict:
        """Process plain text documents."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        extracted_text = file.read()
                    
                    # Basic metadata
                    document_metadata = {
                        'encoding': encoding,
                        'line_count': len(extracted_text.splitlines()),
                    }
                    
                    return {
                        'text': extracted_text,
                        'page_count': 1,  # Assume single page for text files
                        'metadata': document_metadata,
                        'document_type': DocumentType.TXT
                    }
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='replace'
            with open(file_path, 'rb') as file:
                content = file.read()
                extracted_text = content.decode('utf-8', errors='replace')
            
            document_metadata = {
                'encoding': 'utf-8 (with errors replaced)',
                'line_count': len(extracted_text.splitlines()),
            }
            
            return {
                'text': extracted_text,
                'page_count': 1,
                'metadata': document_metadata,
                'document_type': DocumentType.TXT
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {str(e)}")
            raise
    
    def extract_images_from_pdf(self, file_path: Path) -> List[Dict]:
        """
        Extract images from PDF documents for OCR processing.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of dictionaries containing image data and metadata
        """
        try:
            images = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract images from page
                    page_images = page.images
                    for img_idx, img in enumerate(page_images):
                        # Get image data (this is a simplified approach)
                        # In a real implementation, you'd extract the actual image bytes
                        images.append({
                            'page_number': page_num,
                            'image_index': img_idx,
                            'bbox': img.get('bbox', {}),
                            'x0': img.get('x0', 0),
                            'y0': img.get('y0', 0),
                            'x1': img.get('x1', 0),
                            'y1': img.get('y1', 0),
                        })
            
            return images
            
        except Exception as e:
            logger.error(f"Error extracting images from PDF {file_path}: {str(e)}")
            return []
    
    def _get_mime_type_from_extension(self, file_path: Path) -> str:
        """
        Get MIME type from file extension as fallback.
        
        Args:
            file_path: Path to the file
            
        Returns:
            MIME type string
        """
        extension = file_path.suffix.lower()
        
        extension_to_mime = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.txt': 'text/plain',
        }
        
        return extension_to_mime.get(extension, 'application/octet-stream')
    
    def get_document_statistics(self, text: str) -> Dict:
        """
        Calculate basic statistics for extracted text.
        
        Args:
            text: Extracted text content
            
        Returns:
            Dictionary containing text statistics
        """
        if not text:
            return {
                'character_count': 0,
                'word_count': 0,
                'line_count': 0,
                'paragraph_count': 0,
                'sentence_count': 0,
            }
        
        lines = text.splitlines()
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        sentences = text.split('.')
        words = text.split()
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'line_count': len(lines),
            'paragraph_count': len(paragraphs),
            'sentence_count': len([s for s in sentences if s.strip()]),
        }
